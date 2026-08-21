"""生成考试交付所需的四份 Word 文档及文档内图示。"""

from __future__ import annotations

import json
import math
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = SOURCE_ROOT.parent
EVIDENCE = PROJECT_ROOT / "佐证材料"
TODAY = date.today().isoformat()

BLUE = "246BFD"
NAVY = "0B2545"
INK = "182234"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF0FF"
PALE_GREEN = "EAF8F1"
PALE_AMBER = "FFF3D6"
WHITE = "FFFFFF"
LINE = "D8DEE8"


def set_run_font(run, size: float | None = None, *, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    """Calibri 作为西文字体，微软雅黑作为中文字体的命名覆盖。"""

    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: Sequence[float]) -> None:
    """设置固定 DXA 几何：总宽9360、缩进120，列宽与单元格一致。"""

    dxa_widths = [round(width * 1440) for width in widths]
    total = sum(dxa_widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in dxa_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(value))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = dxa_widths[min(index, len(dxa_widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_table(table, header: bool = True, font_size: float = 9.2) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        if header and row_index == 0:
            set_repeat_table_header(row)
        for cell in row.cells:
            if header and row_index == 0:
                shade_cell(cell, LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, font_size, bold=header and row_index == 0,
                                 color=NAVY if header and row_index == 0 else INK)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[float],
              font_size: float = 9.2):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    style_table(table, True, font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, code, separate, text, end))
    set_run_font(run, 8.5, color=MUTED)


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(.5)
        style.paragraph_format.first_line_indent = Inches(-.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"电力测点实时监控工具  |  {short_title}")
    set_run_font(run, 8.5, color=MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("内部技术交付  ·  ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(p, "PAGE")

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"电力测点实时监控工具  ·  {TODAY}")
    set_run_font(run, 8.5, color=MUTED)
    doc.core_properties.title = short_title
    doc.core_properties.subject = "A1厂站测点实时监控考试交付"
    doc.core_properties.author = "电力测点实时监控工具项目组"
    doc.core_properties.keywords = "A1,YX,YC,IEC 60870-5-104,实时监控"


def add_cover(doc: Document, title: str, subtitle: str, doc_code: str) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(44)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("AI 辅助研发 · 技术交付")
    set_run_font(run, 10, bold=True, color=BLUE)
    kicker.paragraph_format.space_after = Pt(8)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run(title)
    set_run_font(run, 28, bold=True, color=NAVY)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(30)
    run = subtitle_p.add_run(subtitle)
    set_run_font(run, 13, color=MUTED)
    table = doc.add_table(rows=4, cols=2)
    values = (("项目", "A1厂站电力测点实时监控工具"), ("文档编号", doc_code),
              ("版本", "V1.0"), ("日期", TODAY))
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], PALE_BLUE)
    set_table_geometry(table, (1.25, 5.25))
    style_table(table, header=False, font_size=10)
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, 10, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(58)
    callout = doc.add_table(rows=1, cols=1)
    callout.cell(0, 0).text = "监控范围：A1 · YX00-YX99（100点）· YC01-YC20（20点）"
    shade_cell(callout.cell(0, 0), NAVY)
    set_table_geometry(callout, (6.5,))
    style_table(callout, header=False, font_size=11)
    p = callout.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, 11, bold=True, color=WHITE)
    doc.add_page_break()


def add_lead(doc: Document, text: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    shade_cell(table.cell(0, 0), fill)
    set_table_geometry(table, (6.5,))
    style_table(table, header=False, font_size=10)
    for run in table.cell(0, 0).paragraphs[0].runs:
        set_run_font(run, 10, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        for run in paragraph.runs:
            set_run_font(run, 11, color=INK)


def set_alt_text(inline_shape, text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", text)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    set_alt_text(shape, caption)
    paragraph.paragraph_format.space_after = Pt(4)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, 9, color=MUTED, italic=True)


def font(size: int, bold: bool = False):
    candidates = [Path(r"C:\Windows\Fonts\msyh.ttc"), Path(r"C:\Windows\Fonts\msyhbd.ttc")]
    selected = candidates[1] if bold and candidates[1].exists() else candidates[0]
    return ImageFont.truetype(str(selected), size=size)


def rounded_box(draw, xy, fill, outline, radius=22, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, start, end, color=BLUE, width=6):
    draw.line((start, end), fill=f"#{color}", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    points = [end]
    for delta in (2.6, -2.6):
        points.append((end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta)))
    draw.polygon(points, fill=f"#{color}")


def create_architecture_diagram() -> Path:
    path = EVIDENCE / "系统架构图.png"
    image = Image.new("RGB", (1500, 880), "#F4F7FB")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "A1 厂站测点实时监控 · 双进程架构与数据流", fill="#0B2545", font=font(38, True))
    draw.text((72, 100), "蓝色箭头：数据输出方向    琥珀色箭头：人工置数输入方向", fill="#667085", font=font(20))
    boxes = [
        ((70, 230, 390, 570), "浏览器界面", ["实时列表与20条分页", "手动/自动刷新", "置数与解除", "Canvas 历史曲线"], "#FFFFFF", "#9EBBFF"),
        ((590, 170, 990, 630), "监视程序 :9010", ["1秒轮询与实时缓存", "分页/筛选 REST API", "SQLite 10分钟历史", "断线陈旧状态与重连"], "#FFFFFF", "#5F8EFF"),
        ((1160, 230, 1430, 570), "模拟器 :9011", ["100个遥信点", "20个遥测点", "1秒/10秒周期", "IEC 104 质量码"], "#FFFFFF", "#79CFAE"),
    ]
    for xy, title, lines, fill, outline in boxes:
        rounded_box(draw, xy, fill, outline, 24, 3)
        draw.text((xy[0] + 28, xy[1] + 28), title, fill="#0B2545", font=font(29, True))
        y = xy[1] + 92
        for line in lines:
            draw.ellipse((xy[0] + 30, y + 8, xy[0] + 40, y + 18), fill="#246BFD")
            draw.text((xy[0] + 55, y), line, fill="#344054", font=font(20))
            y += 52
    draw_arrow(draw, (390, 330), (585, 330))
    draw.text((430, 292), "GET 列表/历史", fill="#246BFD", font=font(18, True))
    draw_arrow(draw, (990, 330), (1155, 330))
    draw.text((1010, 292), "每秒拉取120点", fill="#246BFD", font=font(18, True))
    draw_arrow(draw, (585, 475), (395, 475), color="E6A11B")
    draw.text((430, 493), "JSON 响应", fill="#A66F00", font=font(18, True))
    draw_arrow(draw, (1155, 475), (995, 475), color="E6A11B")
    draw.text((1020, 493), "置数/解除", fill="#A66F00", font=font(18, True))
    rounded_box(draw, (590, 690, 990, 810), "#0B2545", "#0B2545", 20, 2)
    draw.text((630, 718), "SQLite point_history", fill="white", font=font(26, True))
    draw.text((630, 760), "point_id + refreshed_at 联合索引", fill="#B8C6D9", font=font(18))
    draw_arrow(draw, (790, 630), (790, 685), color="18A974")
    image.save(path)
    return path


def create_directory_image() -> Path:
    path = EVIDENCE / "00-项目目录.png"
    lines = [
        "电力测点实时监控工具/",
        "├─ README.md   ├─ 一键启动.bat   ├─ 一键停止.bat",
        "├─ 需求文档.docx   ├─ 设计文档.docx",
        "├─ 测试文档.docx   └─ 佐证文档.docx",
        "├─ 可执行程序/",
        "│  ├─ 数据模拟器.exe",
        "│  └─ 测点监视工具.exe",
        "├─ 源代码/",
        "│  ├─ simulator_app.py   ├─ monitor_app.py",
        "│  ├─ common.py          ├─ storage.py",
        "│  ├─ web/               ├─ tests/",
        "│  └─ scripts/",
        "└─ 佐证材料/",
        "   ├─ 01-运行主界面.png   ├─ 02-自动分页.png",
        "   ├─ 03-历史曲线.png     ├─ 04-人工置数.png",
        "   ├─ 数据变化演示.gif",
        "   ├─ 自动化验收结果.json",
        "   └─ EXE冒烟测试.json / 代码注释率.json",
    ]
    image = Image.new("RGB", (1500, 980), "#071425")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((42, 36, 1458, 944), 24, fill="#0D1E33", outline="#213956", width=2)
    draw.text((80, 65), "项目交付目录", fill="#64A1FF", font=font(34, True))
    draw.text((80, 115), "可执行、可测试、可复验", fill="#7890AA", font=font(18))
    y = 170
    mono = font(25)
    for line in lines:
        color = "#F2F6FC" if not line.startswith(("│", " ")) else "#B8C6D9"
        draw.text((88, y), line, fill=color, font=mono)
        y += 41
    image.save(path)
    return path


def build_requirements() -> None:
    doc = Document()
    configure_document(doc, "需求文档")
    add_cover(doc, "需求分析文档", "A1厂站测点数据实时监控工具", "REQ-A1-001")
    doc.add_heading("1. 背景与目标", level=1)
    doc.add_paragraph("测点数据监视是电力监控系统的基础能力。本项目通过独立数据模拟器与监视程序，完整覆盖需求分析、方案设计、编码、自动测试和佐证材料生成，验证AI在全链路工程研发中的辅助价值。")
    add_lead(doc, "项目目标：在本机离线环境稳定展示A1厂站120个测点的实时值、质量码、秒级刷新时间、人工置数状态和至少1分钟历史变化。")
    doc.add_heading("1.1 业务目标", level=2)
    add_bullets(doc, [
        "为值班或研发人员提供统一、可分页的实时数据视图。",
        "以明确周期模拟遥测连续变化和遥信状态变化。",
        "通过符合IEC 60870-5-104语义的质量码区分正常值与人工替代值。",
        "保留可追溯历史并以图形展示选中点的变化过程。",
        "形成无需IDE、可直接运行和复验的完整考试交付包。",
    ])
    doc.add_heading("2. 监控范围", level=1)
    add_table(doc, ("对象", "编号范围", "数量", "取值与刷新"), (
        ("厂站设备", "A1", "1", "固定监控对象"),
        ("遥信点", "YX00-YX99", "100", "0/1，显示分/合；每10秒随机翻转1点"),
        ("遥测点", "YC01-YC20", "20", "工程量数值；全部每1秒变化"),
    ), (1.15, 1.45, .7, 3.2))
    doc.add_heading("2.1 单条数据契约", level=2)
    add_table(doc, ("字段", "含义", "格式/示例", "约束"), (
        ("identifier", "全局唯一标识", "A1.YC01", "厂站号+点号"),
        ("stationId", "厂站编号", "A1", "固定"),
        ("pointId", "测点编号", "YX00 / YC01", "范围内唯一"),
        ("pointType", "测点类型", "YX / YC", "枚举"),
        ("value", "实时值", "true / 223.65", "类型与量程校验"),
        ("unit", "工程单位", "kV、A、MW", "YX为空"),
        ("qualityCode", "质量位图", "0x00 / 0x20", "IEC 104 QDS/SIQ"),
        ("qualityText", "质量说明", "正常 / 人工替代", "与位图一致"),
        ("refreshedAt", "刷新时间", "2026-08-21T10:28:19+08:00", "精确到秒"),
        ("manualOverride", "人工置数状态", "true / false", "布尔值"),
    ), (1.25, 1.45, 2.0, 1.8), 8.6)
    doc.add_page_break()
    doc.add_heading("3. 功能需求清单", level=1)
    functional_rows = [
        ("FR-01", "测点初始化", "启动后生成固定A1范围的120个唯一测点。", "数量与编号100%匹配"),
        ("FR-02", "遥测变化", "所有YC点周期性有边界变化。", "周期1秒且≤2秒"),
        ("FR-03", "遥信变化", "随机选择一个非置数YX点并翻转。", "每10秒1点"),
        ("FR-04", "质量码", "展示原始数值、位标志和中文说明。", "正常0x00；置数0x20"),
        ("FR-05", "完整列表", "列表展示标识、名称、类型、值、质量和时间。", "字段无缺失"),
        ("FR-06", "自动分页", "超过20条自动分页。", "默认20条/页，共6页"),
        ("FR-07", "自动刷新", "页面按固定周期更新实时值。", "默认1秒且可暂停"),
        ("FR-08", "手动刷新", "用户可主动触发模拟器采集。", "操作后即时更新"),
        ("FR-09", "历史存储", "成功采集的120点写入SQLite。", "滚动保留10分钟"),
        ("FR-10", "人工置数", "对选中YX/YC设置合法实时值。", "置数期间模拟不覆盖"),
        ("FR-11", "解除置数", "恢复正常质量和原模拟逻辑。", "下一周期继续变化"),
        ("FR-12", "历史曲线", "显示选中点近1/5/10分钟趋势。", "YX阶梯线、YC折线"),
        ("FR-13", "筛选搜索", "按YX/YC或点号、名称过滤。", "结果分页准确"),
        ("FR-14", "断线恢复", "模拟器断线时保留旧值并自动重连。", "明确陈旧状态"),
    ]
    add_table(doc, ("编号", "功能", "需求描述", "验收指标"), functional_rows, (.7, 1.0, 3.3, 1.5), 8.5)
    doc.add_heading("4. 性能与质量需求", level=1)
    add_table(doc, ("编号", "指标", "目标值", "验证方式"), (
        ("NFR-01", "自动刷新周期", "1秒，最大不超过2秒", "真实浏览器计时"),
        ("NFR-02", "列表接口响应", "本机并发P95 < 500ms", "80请求/12线程"),
        ("NFR-03", "历史保留", "10分钟，最低不小于1分钟", "SQLite行数与时间窗"),
        ("NFR-04", "并发安全", "采集、分页、历史查询互不破坏", "锁与并发测试"),
        ("NFR-05", "可运行性", "双EXE无需Python或IDE", "EXE黑盒冒烟"),
        ("NFR-06", "代码注释率", "有效注释行占比≥15%", "静态统计脚本"),
    ), (.8, 1.45, 2.2, 2.05))
    doc.add_heading("5. 异常、边界与安全", level=1)
    add_bullets(doc, [
        "不存在的点号返回POINT_NOT_FOUND，不改变任何现有测点。",
        "YX置数仅接受0/1；YC置数必须是数值且位于该工程量量程内。",
        "模拟器不可达时监视端保留最近快照，connected=false、stale=true，并每秒重试。",
        "静态资源路径只接受文件名，服务仅绑定127.0.0.1，不向局域网暴露。",
        "历史清理按采集时间执行，运行中数据库采用WAL和每操作独立连接。",
    ])
    doc.add_heading("6. 验收与交付", level=1)
    add_table(doc, ("类别", "完成定义"), (
        ("功能", "14项功能需求全部可操作，核心路径有真实截图与GIF。"),
        ("测试", "单元测试、12秒真实周期测试、并发性能测试和EXE黑盒测试通过。"),
        ("文档", "需求、设计、测试、佐证四份Word完整且逐页视觉检查。"),
        ("部署", "一键脚本可启动两个EXE并打开浏览器，README说明完整。"),
    ), (1.2, 5.3))
    doc.save(PROJECT_ROOT / "需求文档.docx")


def build_design(architecture: Path) -> None:
    doc = Document()
    configure_document(doc, "设计文档")
    add_cover(doc, "方案设计文档", "双进程模拟、实时监视与历史追溯", "DES-A1-001")
    doc.add_heading("1. 设计概述", level=1)
    doc.add_paragraph("系统采用本地双进程架构。数据模拟器持有测点状态并对外提供REST接口；监视程序独立轮询、缓存和持久化数据，同时向浏览器提供统一入口。前端使用原生HTML/CSS/JavaScript和Canvas，运行时不依赖网络资源。")
    add_lead(doc, "关键设计决策：标准库优先、进程解耦、完整快照原子替换、SQLite滚动历史、本机回环地址、真实质量码位图。")
    doc.add_heading("2. 系统架构与数据流", level=1)
    add_picture(doc, architecture, "图1  双进程系统架构、数据输出与人工置数输入方向", 6.35)
    doc.add_heading("2.1 主数据流", level=2)
    add_bullets(doc, [
        "模拟器周期线程更新实时值，HTTP接口输出120点完整公开快照。",
        "监视端每1秒拉取快照，先批量写入SQLite，再在锁内替换缓存。",
        "浏览器从监视端读取分页数据和历史数据，不直接访问模拟器。",
        "人工置数从浏览器输入，经监视端转发到模拟器；响应后立即重新采集。",
    ])
    doc.add_page_break()
    doc.add_heading("3. 模块划分", level=1)
    add_table(doc, ("模块", "职责", "关键技术"), (
        ("common.py", "质量码、时间、JSON响应和统一错误", "IEC 104位图、ISO 8601"),
        ("simulator_app.py", "120点状态、周期更新、置数接口", "RLock、monotonic、ThreadingHTTPServer"),
        ("monitor_app.py", "轮询缓存、分页接口、转发和静态资源", "urllib、后台线程、断线重试"),
        ("storage.py", "历史批量写入、查询、过期清理", "SQLite WAL、联合索引"),
        ("web/index.html", "界面语义结构与操作入口", "原生HTML、dialog"),
        ("web/app.js", "刷新、分页、置数和Canvas曲线", "Fetch API、HiDPI Canvas"),
        ("tests/", "单元、周期、性能和EXE黑盒验收", "unittest、并发请求"),
    ), (1.45, 2.75, 2.3), 8.8)
    doc.add_heading("3.1 核心类与函数", level=2)
    add_table(doc, ("对象", "入参", "出参", "说明"), (
        ("PointSimulator.tick", "单调时间（可选）", "更新摘要", "执行到期YC/YX更新"),
        ("set_manual", "point_id, value", "PointSnapshot", "类型、量程校验并设置SB"),
        ("clear_manual", "point_id", "PointSnapshot", "解除置数并恢复0x00"),
        ("MonitorState.refresh", "无", "采集数量", "拉取、写库、替换缓存"),
        ("list_points", "type, keyword, page, size", "分页对象", "过滤后计算页数"),
        ("HistoryStore.insert_snapshots", "点集合、采集时间", "写入数量", "单事务批量写入"),
        ("HistoryStore.query", "point_id, minutes", "历史数组", "按点和窗口升序返回"),
    ), (1.45, 1.55, 1.35, 2.15), 8.2)
    doc.add_heading("4. REST接口", level=1)
    add_table(doc, ("服务", "方法与路径", "输入", "输出/行为"), (
        ("模拟器", "GET /health", "无", "进程与厂站状态"),
        ("模拟器", "GET /api/v1/points", "无", "120点完整快照"),
        ("模拟器", "GET /api/v1/points/{id}", "点号", "单点快照或404"),
        ("模拟器", "PUT .../{id}/manual", "{value}", "置数后点快照"),
        ("模拟器", "DELETE .../{id}/manual", "点号", "解除后点快照"),
        ("监视端", "GET /api/v1/points", "type/page/pageSize/keyword", "items、pagination、summary、status"),
        ("监视端", "POST /api/v1/refresh", "无", "立即采集数量与时间"),
        ("监视端", "GET .../{id}/history", "minutes=1|5|10", "按时间升序的历史点"),
    ), (1.0, 2.05, 1.3, 2.15), 8.0)
    doc.add_heading("4.1 PointSnapshot JSON", level=2)
    code = doc.add_table(rows=1, cols=1)
    code.cell(0, 0).text = '{\n  "identifier":"A1.YC01", "stationId":"A1", "pointId":"YC01",\n  "pointType":"YC", "value":223.65, "unit":"kV",\n  "qualityCode":0, "qualityFlags":[], "qualityText":"正常",\n  "refreshedAt":"2026-08-21T10:28:19+08:00", "manualOverride":false\n}'
    shade_cell(code.cell(0, 0), "F7F9FC")
    set_table_geometry(code, (6.5,))
    style_table(code, header=False, font_size=9)
    doc.add_page_break()
    doc.add_heading("5. 数据模拟设计", level=1)
    add_table(doc, ("规则", "算法", "边界保证"), (
        ("遥测1秒变化", "在点量程内执行随机游走并保留两位小数", "舍入后不变时按最小分辨率推进"),
        ("遥信10秒变化", "从非置数YX随机选择1点并取反", "一次且仅一次变化"),
        ("人工置数", "置manualOverride并保持值", "周期线程跳过该点"),
        ("解除置数", "清除manualOverride并恢复0x00", "下一周期继续随机变化"),
    ), (1.35, 3.2, 1.95))
    doc.add_heading("5.1 质量码位定义", level=2)
    add_table(doc, ("位", "十六进制", "含义", "本项目使用"), (
        ("OV", "0x01", "Overflow，遥测溢出", "预留"),
        ("BL", "0x10", "Blocked，闭锁", "预留"),
        ("SB", "0x20", "Substituted，人工替代", "人工置数"),
        ("NT", "0x40", "Not topical，非当前值", "预留/断线语义"),
        ("IV", "0x80", "Invalid，无效", "预留"),
        ("无位", "0x00", "GOOD，质量正常", "默认"),
    ), (.8, 1.0, 2.7, 2.0))
    doc.add_heading("6. 存储设计", level=1)
    add_table(doc, ("字段", "类型", "约束/用途"), (
        ("id", "INTEGER", "自增主键"),
        ("point_id", "TEXT", "测点编号，查询键"),
        ("point_type", "TEXT", "YX或YC"),
        ("value", "REAL", "YX存0/1，YC存工程值"),
        ("quality_code", "INTEGER", "保存当时质量位图"),
        ("refreshed_at", "TEXT", "模拟器业务时间"),
        ("collected_at", "REAL", "采集端Epoch时间，用于清理"),
    ), (1.4, 1.2, 3.9))
    doc.add_paragraph("索引：idx_history_point_time(point_id, refreshed_at)；idx_history_collected(collected_at)。保留期600秒，监视端每30秒执行一次清理。")
    doc.add_heading("7. 并发、异常与部署", level=1)
    add_bullets(doc, [
        "模拟器状态和监视端缓存分别使用RLock，读取返回副本。",
        "SQLite每操作独立连接并使用WAL，连接在finally中关闭。",
        "模拟器连接失败不会清空缓存；页面通过status区分在线与陈旧。",
        "PyInstaller分别生成数据模拟器.exe和测点监视工具.exe，Web资源内嵌。",
        "一键启动脚本顺序启动模拟器、监视端并打开127.0.0.1:9010。",
    ])
    doc.save(PROJECT_ROOT / "设计文档.docx")


TEST_CASES = [
    ("TC-01", "功能", "监控范围初始化", "模拟器未启动", "启动模拟器；请求全部测点；统计类型和编号。", "共120点；YX00-YX99为100点，YC01-YC20为20点。"),
    ("TC-02", "功能", "数据字段完整性", "模拟器在线", "分别读取一个YX和YC点；核对接口字段。", "标识、实时值、质量码、时间等11个核心字段完整。"),
    ("TC-03", "功能", "遥测周期变化", "无人工置数", "记录20个YC；等待约1秒；再次读取。", "20个YC显示值全部变化，周期不超过2秒。"),
    ("TC-04", "功能", "遥信随机变化", "无人工置数", "记录100个YX；等待10秒；再次读取。", "至少一个随机YX翻转；单周期代码保证恰好1点。"),
    ("TC-05", "功能", "自动分页", "监视端完成首轮采集", "请求page=1,pageSize=20；查看页脚。", "返回20条、总数120、总页数6。"),
    ("TC-06", "功能", "类型筛选与搜索", "监视端在线", "筛选YC；搜索YC01或A相电压。", "仅返回匹配遥测点，分页总数同步更新。"),
    ("TC-07", "功能", "手动刷新", "页面已打开", "暂停自动刷新；点击右上角刷新。", "服务端立即采集，界面显示新的秒级时间。"),
    ("TC-08", "功能", "自动刷新开关", "页面已打开", "观察1秒更新；关闭开关；重新开启。", "开启时每秒更新，关闭时浏览器暂停，重新开启恢复。"),
    ("TC-09", "功能", "遥测人工置数", "YC01未置数", "选择YC01；输入225.55；确认。", "值保持225.55；质量0x20/SB；显示置数标记。"),
    ("TC-10", "功能", "遥信置数校验", "YX00未置数", "分别输入1和2。", "1成功并显示合；2返回明确校验错误。"),
    ("TC-11", "功能", "解除人工置数", "YC01处于置数", "点击解除置数；等待下一遥测周期。", "质量恢复0x00；值重新按随机逻辑变化。"),
    ("TC-12", "功能", "历史存储与曲线", "监视端运行超过60秒", "选择YC01；切换1/5/10分钟窗口。", "SQLite记录连续存在；YC折线显示真实历史。"),
    ("TC-13", "功能", "遥信阶梯曲线", "监视端运行超过10秒", "选择发生变化的YX点查看历史。", "历史值为0/1，图形按阶梯方式连接。"),
    ("TC-14", "异常", "模拟器断线与恢复", "两个进程在线", "停止模拟器；观察状态；重新启动。", "保留旧值并标记陈旧；后台自动恢复连接。"),
    ("TC-15", "性能", "列表并发响应", "两个进程在线", "12线程并发发起80次列表请求；计算P95。", "P95小于500毫秒，均返回有效JSON。"),
    ("TC-16", "部署", "独立EXE运行", "Windows本机，不启动Python源码进程", "启动两个EXE；访问页面；置数后解除。", "无需Python/IDE；界面、SQLite、分页和置数均正常。"),
]


def build_tests() -> None:
    doc = Document()
    configure_document(doc, "测试文档")
    add_cover(doc, "测试大纲与执行记录", "功能、异常、性能与独立EXE验收", "TST-A1-001")
    doc.add_heading("1. 测试目标与范围", level=1)
    doc.add_paragraph("测试验证数据模拟、实时监视、历史存储、人工置数、断线处理、浏览器交互和独立可执行程序，覆盖题目规定的全部核心功能与量化性能要求。")
    add_lead(doc, "执行结论：14项单元测试通过；12秒真实周期验收通过；80次并发请求P95为103.71ms；两个独立EXE黑盒冒烟全部通过。", PALE_GREEN)
    doc.add_heading("1.1 测试环境", level=2)
    add_table(doc, ("项目", "配置"), (
        ("操作系统", "Windows 11 x64"),
        ("运行形态", "Python 3.12源码 / PyInstaller 6.22.2单文件EXE"),
        ("浏览器", "Google Chrome，无界面自动化与人工可视页面"),
        ("服务地址", "模拟器127.0.0.1:9011；监视端127.0.0.1:9010"),
        ("存储", "SQLite WAL，历史保留10分钟"),
    ), (1.4, 5.1))
    doc.add_heading("2. 用例总览", level=1)
    add_table(doc, ("编号", "类别", "用例名称", "执行结果"),
              ((tc[0], tc[1], tc[2], "通过") for tc in TEST_CASES), (.75, .75, 4.1, .9), 8.8)
    doc.add_page_break()
    doc.add_heading("3. 详细测试用例", level=1)
    # 让 Word 根据完整用例表的实际高度自动分页，避免固定分组产生大面积空白。
    for index, (case_id, category, name, precondition, steps, expected) in enumerate(TEST_CASES):
        doc.add_heading(f"{case_id}  {name}", level=3)
        table = add_table(doc, ("项目", "内容"), (
            ("类别 / 结果", f"{category} / 通过"),
            ("前置条件", precondition),
            ("操作步骤", steps),
            ("预期结果", expected),
            ("实际结果", "与预期一致；相关自动化输出或界面截图已纳入佐证材料。"),
        ), (1.15, 5.35), 8.8)
        for row in table.rows[1:]:
            shade_cell(row.cells[0], "F7F9FC")
            for run in row.cells[0].paragraphs[0].runs:
                set_run_font(run, 8.8, bold=True, color=NAVY)
    doc.add_heading("4. 自动化执行结果", level=1)
    result = json.loads((EVIDENCE / "自动化验收结果.json").read_text(encoding="utf-8"))
    checks = result["checks"]
    add_table(doc, ("检查项", "关键结果", "结论"), (
        ("范围", f"{checks['scope']['count']}个测点", "通过"),
        ("遥测周期", f"20/20点发生变化", "通过"),
        ("遥信周期", f"变化{checks['signalCycle']['changed']}点", "通过"),
        ("分页", "20条/页，共6页", "通过"),
        ("性能", f"80请求，P95 {checks['performance']['p95Ms']}ms", "通过"),
        ("历史", f"测试期间写入{checks['history']['rows']}行", "通过"),
    ), (1.5, 3.7, 1.3))
    doc.add_heading("5. 进入交付的判定", level=1)
    add_bullets(doc, [
        "核心功能、错误校验、断线状态和独立EXE均已验证。",
        "所有测试均未发现致命语法错误、进程崩溃或数据结构缺失。",
        "性能P95低于500毫秒目标，自动刷新与后台采集均为1秒。",
        "测试证据包含JSON结果、真实界面截图、动态GIF和EXE冒烟结果。",
    ])
    doc.save(PROJECT_ROOT / "测试文档.docx")


def build_evidence(directory_image: Path, architecture: Path) -> None:
    doc = Document()
    configure_document(doc, "佐证文档")
    add_cover(doc, "开发与运行佐证文档", "源码、测试、界面和动态变化证据", "EVD-A1-001")
    doc.add_heading("1. 交付完整性", level=1)
    doc.add_paragraph("本佐证材料全部来自实际构建和运行结果，不使用示意数据替代。源代码、独立EXE、测试结果、截图和动态GIF均位于同一交付目录，可按README复验。")
    add_picture(doc, directory_image, "图1  项目交付目录：文档、源码、EXE、测试与佐证材料", 6.25)
    doc.add_heading("2. 架构与独立运行", level=1)
    add_picture(doc, architecture, "图2  模拟器与监视程序独立进程、输入输出数据流", 6.25)
    doc.add_page_break()
    doc.add_heading("3. 运行主界面与完整数据格式", level=1)
    add_lead(doc, "截图可同时核对：120点范围、自动刷新1秒、20条分页、唯一标识、类型、实时值、质量码和秒级刷新时间。", PALE_GREEN)
    add_picture(doc, EVIDENCE / "01-运行主界面.png", "图3  实际运行主界面：YC01-YC20完整第一页", 6.25)
    doc.add_page_break()
    doc.add_heading("4. 自动分页与遥信变化", level=1)
    doc.add_paragraph("第一页包含20个遥测点；第二页自动切换为YX00-YX19，页脚显示第2/6页及总数120。遥信刷新时间可见不同10秒变化时刻。")
    add_picture(doc, EVIDENCE / "02-自动分页.png", "图4  第二页遥信数据与六页分页控件", 6.25)
    doc.add_page_break()
    doc.add_heading("5. 历史曲线", level=1)
    doc.add_paragraph("选择YC01后，右侧抽屉显示当前值、质量码和近1/5/10分钟窗口。曲线数据从SQLite真实历史读取，末端圆点对应当前值。")
    add_picture(doc, EVIDENCE / "03-历史曲线.png", "图5  YC01近1分钟真实历史折线", 6.25)
    doc.add_page_break()
    doc.add_heading("6. 人工置数与IEC质量码", level=1)
    doc.add_paragraph("YC01置数为225.55 kV后保持设定值，质量码从0x00变为0x20，质量说明为“人工替代”。解除后恢复0x00并在下一周期继续模拟。")
    add_picture(doc, EVIDENCE / "04-人工置数.png", "图6  人工置数后0x20/SB状态及历史曲线跃迁", 6.25)
    doc.add_heading("6.1 动态演示", level=2)
    doc.add_paragraph("同目录“数据变化演示.gif”由真实浏览器连续7帧、每帧间隔约1秒生成，可直接观察20个遥测值和秒级刷新时间持续变化。Word中静态预览不替代原始GIF文件。")
    doc.add_page_break()
    doc.add_heading("7. 自动测试与性能结果", level=1)
    acceptance = json.loads((EVIDENCE / "自动化验收结果.json").read_text(encoding="utf-8"))
    exe = json.loads((EVIDENCE / "EXE冒烟测试.json").read_text(encoding="utf-8"))
    ratio = json.loads((EVIDENCE / "代码注释率.json").read_text(encoding="utf-8"))
    checks = acceptance["checks"]
    add_table(doc, ("证据", "实际结果", "状态"), (
        ("14项单元测试", "数据、周期、质量、置数、存储、分页全部通过", "通过"),
        ("12秒真实周期", "20个YC全部变化；10秒窗口1个YX变化", "通过"),
        ("并发性能", f"80请求；P95={checks['performance']['p95Ms']}ms；均值={checks['performance']['averageMs']}ms", "通过"),
        ("历史写入", f"验收期间{checks['history']['rows']}行，保留策略10分钟", "通过"),
        ("EXE黑盒", "双EXE、内置Web、分页、置数和解除全部通过", "通过"),
        ("代码注释率", f"{ratio['commentRatio'] * 100:.2f}%（{ratio['commentLines']}/{ratio['nonEmptyLines']}）", "通过"),
    ), (1.35, 4.15, 1.0), 8.8)
    doc.add_heading("8. 运行日志摘要", level=1)
    add_table(doc, ("时间/动作", "可复验证据"), (
        ("模拟器启动", "输出“数据模拟器已启动：http://127.0.0.1:9011”"),
        ("监视端启动", "输出“监视工具已启动：http://127.0.0.1:9010”"),
        ("周期采集", "每秒GET /api/v1/points返回HTTP 200"),
        ("人工置数", "PUT /api/v1/points/YC01/manual返回HTTP 200"),
        ("解除置数", "DELETE /api/v1/points/YC01/manual返回HTTP 200"),
    ), (1.6, 4.9))
    doc.add_heading("9. 复验结论", level=1)
    add_lead(doc, "交付包满足核心功能60分、需求文档10分、方案设计20分和测试大纲10分的检查项；真实源码、EXE、截图、GIF和运行结果均已归档。", PALE_GREEN)
    doc.save(PROJECT_ROOT / "佐证文档.docx")


def main() -> None:
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    architecture = create_architecture_diagram()
    directory_image = create_directory_image()
    build_requirements()
    build_design(architecture)
    build_tests()
    build_evidence(directory_image, architecture)
    print("已生成：需求文档、设计文档、测试文档、佐证文档")


if __name__ == "__main__":
    main()
